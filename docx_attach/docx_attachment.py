import base64
import os
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Optional
from copy import deepcopy

from docx import Document
from docx.opc.constants import CONTENT_TYPE, RELATIONSHIP_TYPE
from docx.opc.oxml import parse_xml
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from PIL import Image, ImageDraw, ImageFont


@dataclass
class AttachmentType:
    """附件类型的数据类"""
    content_type: str
    file_name_prefix: str
    file_extension: str
    program_id: str
    shape_width: int
    shape_height: int
    icon_base64: str


# 预定义的附件类型
XLSX_ATTACHMENT = AttachmentType(
    content_type=CONTENT_TYPE.SML_SHEET,
    file_name_prefix="Microsoft_Excel_",
    file_extension="xlsx",
    program_id="Excel.Sheet.12",
    shape_width=76,
    shape_height=60,
    icon_base64=(
        "iVBORw0KGgoAAAANSUhEUgAAAEAAAABACAYAAACqaXHeAAAAAXNSR0IArs4c6QAABXFJREFUeF7tW3tQVFUc/u7dxQU"
        "1bPJZoBLQDIWwV0EqnRo0iLKEQEbyQVpjCgnGkDkRwWypNWpOBiNPS0GzUEHy1eAjBcEnfzQp5oROWWg6IiEUk+De25"
        "xloL3cu487tevZ7v7+2/O49/d953e/3zlnz2GgcmNowx9qCJ/wvaHxvLP8umcERBoiPdvQwTFGgeMFhmNY6CGAA+Dpp"
        "fEYdtpwusMZJDiFgDBD2IjuuzzHsqyeFwSOIUAZTLAE0KUJ4Aycn5HX6hme5xgGnADTqPopGU2XIUCfow/mGU3viPaO"
        "qh7ASCVg5dpSR0BgXqBu8M1hHCMInMAwHCAQoGRkvf4tWOoICMqKGD7Io4eMqt4EtleYQhwBlBoN4HK4yUawWQwEDmA"
        "ediZYuXf5hPrka7Wabkf5cRco/Ca++jJ5vikLEAJ4sGcc9UKlz/XV+4LVsEq7KWq/76VqFgwE1RLA8GzE3llVZ90EqP"
        "UTcEeA+xNwa4BbBKnNAiey6zBEN0SU12dWx/f/nhOUhLlBL9ud99+tz8G5VvEWA9Ui6CaAtghInbZENtwECCg6WmIxF"
        "CP8JyNs/CTZ+sKjxRb7URcBAaP8UZW2U+Jwj7EH4e8/YRFI+euboR8bKqr/804XlpSl4lyL5W0/6gggCFYlfICZ3AsS"
        "sBsO5mFzfZmkPCr4GaxPWispzz9SgE21n1kVMGcshhSLoE6rw5ncExLHr9++gZj1MyTlX6Zsw2MPPSoqP/tTIxaXpYL"
        "nedcjgHhMtCBl2mKJ82sOrMP2U1/1l88Kj0du7HuSdqnlaThx6aTN9EVlBPR53ZBdi6G6oSIQTVcvYG5xcn+Z3OhvaS"
        "jHJzWf2gRPGpSmFMFrkHi3jeTyPiPzgJARFjeVpdF4seK/mwcQHSB6MNBW7MhCzfmDSAxPQE5stqi6vasdcXkJaO+6b"
        "RcBVIqguecVqV8g6MEgEZi6H48jfVsG5JS/jxy70AOgnoCJ4zhsWSRV8jlF80HC39zqmxuwdOsye7Gb2lFPAHHy46Q1"
        "iA6OEgHb1Vhl+gTM7cUNcfi1reX/R8DI+0bg8Ns1VoEVHytFwbdFisCTxjHTn4VWqxX1M1/MKBHAvoc4ZDGUEZ2OV59"
        "aKAvwZmcrotbFKAZPOlCdBs0RLX8uE8lT5lkESTThwrUfFJPgEgQEjgpAZdoOq+AIeEKCUnOJT2Db4jKE+IonIycvn8"
        "KTAeLF0co9q0HEUYlRnwXkprpXbv2Csvpy5MZJp8DT10bj1h9tdnNANQFjho1BzVv7JWCI2lc27sb2lHKM9h4tSY8kE"
        "uw1qglYGW9A7MSZIizX2n/D7II56PyrE4uefg3pUUslWN/Ymo6GZulqUo4UagkgEx8yARpoGw7lYfPx3j0BDavB18uq"
        "MPYBX1EzIojzSxbAyBttBgKViyFvL28Q4Rs/fJwIwM+tV0wLHXMjqZGkyIGWf3gjNtV9bpMAKtNgZkwGFkz9Z8nbh8K"
        "Syu9O3wn/kf4SsPOKX8H5q01WSaCOgKmPTEFBcr7E6eYbl5C4MUkWzIzQ5/FR4ipJ3aGmI1hescJ1CPDQeKBkYSEmjZ"
        "8ocfrN7Zk4drHWIhi5uQJpvHLPh9jVWGmxH1UiaGlLnHhvbWub1D/uHyFLHKnb+91+tPwuv0qkigCbiuWABm4CaPtny"
        "AGDbPWRqo8A6tKgsyPATYATzgkq/mvMmVHgjgB3BDj+qKz7E1D7OcF98dWmY8L9Z4WpOi4f7LOa9dDccZTwCsbu0gOz"
        "D1zvJ8D8Raq8MGGL6cD0QN3g+1VyZcYWGeb1qrk0pYQU7h3OzzhIJdfm7CVGVRcn7SXFzxDp6d3TwTGsCq7O2ksKaae"
        "ay9NKSHFk278BpaKbbiHMt1QAAAAASUVORK5CYII="
    ),
)

DOCX_ATTACHMENT = AttachmentType(
    content_type=CONTENT_TYPE.WML_DOCUMENT,
    file_name_prefix="Microsoft_Word_",
    file_extension="docx",
    program_id="Word.Document.12",
    shape_width=76,
    shape_height=60,
    icon_base64=(
        "iVBORw0KGgoAAAANSUhEUgAAAEAAAABACAYAAACqaXHeAAAAAXNSR0IArs4c6QAABv5JREFUeF7tWwtQVFUY/i7vl5K"
        "VqfnAGHWXx6KA7ydkCSoiaIsgoYapoDA2YySBguazTGc0hEghEwUfmKhJYFmIjzE1lRFIMHzjIwSZTDNl9zTnLnu5S9"
        "I+WOAC+8/s7L3nnnvO/3/nP99/XpdBM4uTNLmXHJAYARJiBAkIJAD744QQJu+3vXM9m0M1pqkqsZcm21oCEtZYI0iIw"
        "lAxgFfV1dnqABBLkyXGyhYF4wAiFwOMSJ2hDT0XLABK9wUDJwYQKYwk1FC1raoNGC0OAOu+jNyZwFjEgIgIw4ga26qC"
        "BUAkTRGZQCaSM0TclK0qWAAcA5KJNsq1kbzXAZwt3jMvgGmnANS2I/lckAAMcnwdC6RuTepss5Z/R8uvEiwA2+J9mgy"
        "As8V3UAsADAAIkQNoFzB4QHvvAoMdu+mNA+ZL3VXKEjwH6M1yAC/qTgYAWkMU0Oc4gHoBXwTvAYYoYAiDbWgcQN3Z1s"
        "Yc1pamsLKgPxOU3a5G3q839En2DZalcxSwsTSDjZUp6L+1lRkKSu+rVNK9cwesjfTESzYWrIH0F5+cj6xjpSr5/D1EW"
        "Bk+hkt7UP0EgTFZuFv5l7ABKNo9l1Ow9GYV/KMyVRQe7doLSdHeKmmphwqwfscvKmmR0wYibErdzO5EwS3MW/19sxhP"
        "K9HZA76O88FgJ0X4ePqsBt6Ru1BR/YRTPHRSfyx6d4iKIfkXbiJ8bY5K2qeRb8JnZB8uLeXARWxIP/O/AAgiCoRPdUN"
        "EwEBOUTp9pPFTKavCPeDn0U/FkPKKR5iwcDdqZHIuPWOlH1z6vsbdR206iuyTZcIHwF3cFduX+3KK0v6d+dNl7n7PGn"
        "842Xf+jyEBH3+LoqsP2HQLMxOcSpkJczNj9v7ZcxmmLt6Hq+XVwgeAkhpVXikpBwuwYaeif1MCzN44DSbGRvj7nxpYm"
        "ptw+WIT8zgi7NOjEw6sl3LPzpfcQ0jcQbX9X98rQjqPBAvS32eNpPLjmWtYuP4H9ppPgBdK7sNV1IUzik+EHu522PyR"
        "F/csPbcIq1JPqgVAnxl0JkGqxJbYCRju0oPVhx8JQn37Y1GwggB35hQi2NuZ05lPhCETnBE9czj3LC45H/t43UifhjZ"
        "UVqMAmOM3AB8EDWbLppHAPSSVveYze0xiHmZOlEBk9wr7jBLhuIgM9npF2BhM8azbIZsWsx+FZRVq7a7vsmpfUJOh/u"
        "qSxpMhDzc7bF5c58JvR2TgTsUjZK17B317vcxW6/dhJguAP89Q30V7UXb7IfhEea28Gn5RmSoRQptWaywI/Pc1BsCuq"
        "y1LdkqJWJeL0htVOJIQxCZRAhw4IxVBXk5YEjqCyxe18SiyT5WBzyGHjl9BdMLPGtkhiHGAUtPzabO5MPbF7nMouVmJ"
        "hCiFV1B3pm5N4zyN90rZknURadmXkP9VCJe29ptTSMsubH0AJEWPx2jXnqziuaevsh5Ah7dUsvJKEZuUB2tLM5zcGgJ"
        "TE0W8p0S4/fAlbF0ykTN41vJDOFt8V2MABLMmSEmQkiEVGgmu3amG11B79p6O+2nYo7JzxWQM6KcIh5QI03OKEBUylL"
        "1/+OdTjF+4C4+ePNMIAH1malQUoIpMHtMPq+d7sDrRSHCv8jF6d7Nl72evOIzTheXsdcx7IxDs7cTpfjD/CnxH92Xvj"
        "1+8hbA1zTcB4gPYaADq929l4ZQAPcN2cK0qHeuAZXNHcXVfvl4JcW9FaEzadx4Je85p3LBUacF0gQ5WZjiaGMwuaPBF"
        "SYDKNDdRV6R9Ujd3IARgak8gLfgsV6sFEEFFAWrgtvhJGFRvo0JJgEoA6NwhZ1MgOlqbqwBFCMHY+em4X/VYKw8Q1NZ"
        "Y/JxRCHjLQcUAPgEqH6Qu9cEQZ9Ul6Eu//4HA2CyNjacZBecBMyZKsHjGMBUj6MIHDXd8iQ0dgeledURIn2UcKcbKlB"
        "NaA6DVC2oy6zwUVpY7yrUnhjl3h20HC1BOoHwQ92U+G+74EjjOEUtnj1RJW5J0DPvzSvRpj1ZlNToKaFMbfxLT0dqMB"
        "av+Iqk25ekjb7MCoA+FlRxg2Bpr7+cDBBUG9eXampYjmDVBTRUWcr5WSYL6BLRVAiCYNUF9toQ2ZQluKKyN8vrIawDA"
        "cEKkDZ0Q0aVLGLpAMx2XJ0CFIE+L6+I1urxDQBLbMQDy5XI58hgXaYqoxui5mCGMiABiwkDEEPYDR8XmXwtIi382J5m"
        "e2EkmN3VmiFxECBTgMISCUnfYpwmBaXEAGrJNMiXRXmZqKgGROzGgHiMXK/6h2CnRkwgWgBfZ5yhdZiZHd4kRiIQwxI"
        "UBHOj3hwR4Q1c8WhUADRnpHLSli0xGOGBAGPrxND0tYaMOmDYBQIPdaGqyWGai+JqcIcQF7Cf0jGK3tVaaE4B/Ac9my"
        "WiNuu1sAAAAAElFTkSuQmCC"
    ),
)


class AttachmentHandler:
    """处理Word文档中的附件嵌入"""

    def __init__(self, doc: Document, file_path: str, attachment_type: AttachmentType):
        self.doc = doc
        self.file_path = Path(file_path)
        self.attachment_type = attachment_type
        self.doc_part = doc.part
        self.package = doc.part.package

        self.file_content = self._read_file_content()
        self.shape_rid: Optional[str] = None
        self.embed_rid: Optional[str] = None

    def _read_file_content(self) -> bytes:
        """读取附件文件内容"""
        return self.file_path.read_bytes()

    def replace_placeholder(self, placeholder: str) -> None:
        """替换文档中的占位符为附件"""
        self._create_relationships()
        object_element = self._build_object_element()
        self._replace_in_document(placeholder, object_element)

    def _create_relationships(self) -> None:
        """创建所需的关系"""
        self._create_shape_relationship()
        self._create_embedded_relationship()

    def _create_shape_relationship(self) -> None:
        """创建图标关系"""
        shape_name = f"icon_{self.attachment_type.file_extension}_{len(self.package.parts)}.png"
        shape_path = f"/word/media/{shape_name}"

        shape_part = Part(
            partname=PackURI(shape_path),
            content_type=CONTENT_TYPE.PNG,
            blob=self._generate_icon(),
            package=self.package
        )

        self.package.parts.append(shape_part)
        self.shape_rid = self.doc_part.relate_to(shape_part, RELATIONSHIP_TYPE.IMAGE)

    def _create_embedded_relationship(self) -> None:
        """创建嵌入文件关系"""
        file_name = (
            f"{self.attachment_type.file_name_prefix}{len(self.package.parts)}."
            f"{self.attachment_type.file_extension}"
        )
        embed_path = f"/word/embeddings/{file_name}"

        embed_part = Part(
            partname=PackURI(embed_path),
            content_type=self.attachment_type.content_type,
            blob=self.file_content,
            package=self.package
        )

        self.package.parts.append(embed_part)
        self.embed_rid = self.doc_part.relate_to(embed_part, RELATIONSHIP_TYPE.PACKAGE)

    def _generate_icon(self) -> bytes:
        """生成附件图标"""
        title = self.file_path.name
        title2 = ''
        name, ext = os.path.splitext(title)

        # 处理文件名长度
        if len(name) >= 14:
            title = name[:14]
            title2 = name[14:20] + ext
        else:
            title = name + ext

        # 创建基础图像
        base_image = Image.new(
            'RGBA',
            (self.attachment_type.shape_width, self.attachment_type.shape_height),
            (255, 255, 255, 0)
        )

        # 添加图标
        icon_image = Image.open(BytesIO(base64.b64decode(self.attachment_type.icon_base64)))
        icon_image.thumbnail((30, 30))

        # 设置字体
        try:
            font = ImageFont.truetype('Arial', 8)
        except IOError:
            font = ImageFont.load_default()

        # 绘制文本
        draw = ImageDraw.Draw(base_image)

        # 绘制第一行文本
        font_length = font.getlength(title)
        x_position = (base_image.width - font_length) / 2
        draw.text((int(x_position), 32), title, font=font, fill=(0, 0, 0, 255))

        # 绘制第二行文本（如果有）
        if title2:
            font_length2 = font.getlength(title2)
            x_position2 = (base_image.width - font_length2) / 2
            draw.text((int(x_position2), 40), title2, font=font, fill=(0, 0, 0, 255))

        # 粘贴图标
        icon_position = ((base_image.width - icon_image.width) // 2, 0)
        base_image.paste(icon_image, icon_position, icon_image)

        # 保存为PNG
        output = BytesIO()
        base_image.save(output, format='PNG')
        return output.getvalue()

    def _build_object_element(self) -> object:
        """构建对象元素XML"""
        shape_id = f"_x0000_i{os.urandom(4).hex()}"

        object_xml = self._get_object_xml_template().format(
            shape_id=shape_id,
            shape_width=self.attachment_type.shape_width,
            shape_height=self.attachment_type.shape_height,
            shape_rid=self.shape_rid,
            program_id=self.attachment_type.program_id,
            embed_rid=self.embed_rid
        )

        return parse_xml(object_xml)

    @staticmethod
    def _get_object_xml_template() -> str:
        """获取对象XML模板"""
        return (
            '<w:object xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'xmlns:v="urn:schemas-microsoft-com:vml" '
            'xmlns:o="urn:schemas-microsoft-com:office:office" '
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            '<v:shape id="{shape_id}" type="#_x0000_t75" style="width:{shape_width}pt;height:{shape_height}pt">'
            '<v:imagedata r:id="{shape_rid}" o:title=""/>'
            '</v:shape>'
            '<o:OLEObject Type="Embed" ProgID="{program_id}" ShapeID="{shape_id}" '
            'DrawAspect="Icon" ObjectID="_1" r:id="{embed_rid}"/>'
            '</w:object>'
        )

    def _replace_in_document(self, placeholder: str, object_element) -> None:
        """在文档中替换占位符"""
        placeholder_text = f"{{{placeholder}}}"

        for paragraph in self.doc.paragraphs:
            if placeholder_text not in paragraph.text:
                continue

            self._replace_in_paragraph(paragraph, placeholder_text, object_element)

    def _replace_in_paragraph(self, paragraph, placeholder_text: str, object_element) -> None:
        """在段落中替换占位符"""
        # 处理单个run包含完整占位符的情况
        if self._replace_in_single_run(paragraph, placeholder_text, object_element):
            return

        # 处理占位符跨多个run的情况
        self._replace_across_runs(paragraph, placeholder_text, object_element)

    def _replace_in_single_run(self, paragraph, placeholder_text: str, object_element) -> bool:
        """处理单个run包含完整占位符的情况"""
        for i, run in enumerate(paragraph.runs):
            if placeholder_text not in run.text:
                continue

            prefix = run.text[:run.text.index(placeholder_text)]
            postfix = run.text[run.text.index(placeholder_text) + len(placeholder_text):]

            run.text = prefix
            postfix_run = paragraph.add_run(postfix)
            self._copy_run_format_xml(postfix_run, run)

            object_run = paragraph.add_run()
            object_run.element.append(object_element)

            # 按顺序插入新的runs
            paragraph._p.insert(i + 1, object_run.element)
            paragraph._p.insert(i + 2, postfix_run.element)
            return True

        return False

    def _replace_across_runs(self, paragraph, placeholder_text: str, object_element) -> None:
        """处理占位符跨多个run的情况"""
        start_run_index = None
        first_run_prefix_index = -1
        accumulated_text = ""
        runs_to_clear = []

        for i, run in enumerate(paragraph.runs):
            if first_run_prefix_index == -1:
                first_run_prefix_index = self._part_contains(run.text, placeholder_text)
            if start_run_index is None and first_run_prefix_index != -1:
                start_run_index = i

            if start_run_index is not None:
                runs_to_clear.append(run)
                accumulated_text += run.text

                if self._process_accumulated_text(
                        paragraph,
                        placeholder_text,
                        accumulated_text,
                        runs_to_clear,
                        start_run_index,
                        first_run_prefix_index,
                        object_element
                ):
                    break
                elif not placeholder_text.startswith(accumulated_text):
                    start_run_index = None
                    first_run_prefix_index = -1
                    accumulated_text = ""
                    runs_to_clear = []

    @staticmethod
    def _process_accumulated_text(
            paragraph,
            placeholder_text: str,
            accumulated_text: str,
            runs_to_clear: list,
            start_run_index: int,
            first_run_prefix_index: int,
            object_element
    ) -> bool:
        """处理累积的文本"""
        if placeholder_text not in accumulated_text:
            return False

        first_run = runs_to_clear[0] if runs_to_clear else None
        last_run = runs_to_clear[-1] if runs_to_clear else None
        if first_run is None or last_run is None:
            return False

        # 清除中间的runs
        if len(runs_to_clear) > 1:
            for run_to_clear in runs_to_clear[1:len(runs_to_clear) - 1]:
                run_to_clear.text = ""

        # 处理第一个run
        first_run.text = first_run.text[:first_run_prefix_index]

        # 处理最后一个run
        for j in range(1, len(placeholder_text) + 1):
            postfix = placeholder_text[j:]
            if last_run.text.startswith(postfix):
                last_run.text = last_run.text[len(postfix):]
                break

        # 插入对象
        object_run = paragraph.add_run()
        object_run.element.append(object_element)
        paragraph._p.insert(start_run_index + 1, object_run.element)
        return True

    @staticmethod
    def _copy_run_format_xml(source_run, target_run):
        # 复制rPr(run properties)元素
        source_rPr = source_run.element.rPr
        if source_rPr is not None:
            target_run.element.get_or_add_rPr().clear_content()
            for element in source_rPr.iterchildren():
                target_run.element.rPr.append(deepcopy(element))

    @staticmethod
    def _part_contains(run_text: str, placeholder_text: str) -> int:
        index = -1

        if placeholder_text.startswith(run_text):
            return 0
        for i in range(1, len(placeholder_text) + 1):
            prefix = placeholder_text[:i]
            if run_text.endswith(prefix):
                index = run_text.index(prefix)
            else:
                break
        return index


def replace_attachment_in_document(
        doc: Document,
        placeholder: str,
        file_path: str,
        attachment_type: AttachmentType
) -> None:
    """替换文档中的占位符为附件"""
    handler = AttachmentHandler(doc, file_path, attachment_type)
    handler.replace_placeholder(placeholder)


def replace_attachment_in_template(
        template_path: str,
        output_path: str,
        placeholder: str,
        file_path: str,
        attachment_type: AttachmentType
) -> None:
    """在模板中替换占位符并保存为新文档"""
    doc = Document(template_path)
    replace_attachment_in_document(doc, placeholder, file_path, attachment_type)
    doc.save(output_path)


# 便捷函数
def replace_xlsx(doc: Document, placeholder: str, file_path: str) -> None:
    """替换为Excel附件"""
    replace_attachment_in_document(doc, placeholder, file_path, XLSX_ATTACHMENT)


def replace_docx(doc: Document, placeholder: str, file_path: str) -> None:
    """替换为Word附件"""
    replace_attachment_in_document(doc, placeholder, file_path, DOCX_ATTACHMENT)


def replace_xlsx_in_template(
        template_path: str,
        output_path: str,
        placeholder: str,
        file_path: str
) -> None:
    """在模板中替换为Excel附件"""
    replace_attachment_in_template(
        template_path,
        output_path,
        placeholder,
        file_path,
        XLSX_ATTACHMENT
    )


def replace_docx_in_template(
        template_path: str,
        output_path: str,
        placeholder: str,
        file_path: str
) -> None:
    """在模板中替换为Word附件"""
    replace_attachment_in_template(
        template_path,
        output_path,
        placeholder,
        file_path,
        DOCX_ATTACHMENT
    )